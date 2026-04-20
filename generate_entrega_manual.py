from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()

    # Portada Profesional
    title = doc.add_heading('PROYECTO FINAL: UNIDADES 3 Y 4\nAPLICACIÓN MÓVIL FASTBITE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n' * 3)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('ALUMNO: Henrique\nMATRÍCULA: 2302070\nFECHA: 14 de Abril de 2026')
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph('\n' * 2)
    repo = doc.add_paragraph()
    repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    repo.add_run('URL GITHUB:\nhttps://github.com/2302070-commits/Proyecto_Unidad_4_FastBite').font.color.rgb = RGBColor(0, 0, 255)

    doc.add_page_break()

    # Índice
    doc.add_heading('ÍNDICE', level=1)
    doc.add_paragraph('1. Introducción\n2. Problemática Detectada\n3. Propuesta de Solución\n4. Colores en Hexadecimal\n5. Diagrama de Base de Datos\n6. Tecnologías y Navegación\n7. Guía de Creación de APK\n8. Conclusión')

    doc.add_page_break()

    # 1-3. Secciones generales
    doc.add_heading('1. INTRODUCCIÓN', level=1)
    doc.add_paragraph('FastBite es una aplicación híbrida desarrollada en React Native. Su propósito es optimizar el consumo de alimentos mediante el rescate de excedentes y el soporte nutricional personalizado.')

    doc.add_heading('2. PROBLEMÁTICA DETECTADA', level=1)
    doc.add_paragraph('Desperdicio masivo de alimentos en locales comerciales y falta de herramientas interactivas para dietas especializadas en el idioma local.')

    doc.add_heading('3. PROPUESTA DE SOLUCIÓN', level=1)
    doc.add_paragraph('Una plataforma centralizada que integra:')
    doc.add_paragraph('Rescate de platillos con descuento.', style='List Bullet')
    doc.add_paragraph('Filtros para dietas saludables.', style='List Bullet')
    doc.add_paragraph('Recetario con traducción en tiempo real.', style='List Bullet')

    # 4. COLORES EN HEXADECIMAL (REQUISITO)
    doc.add_heading('4. COLORES EN HEXADECIMAL', level=1)
    doc.add_paragraph('Se ha definido una paleta cromática profesional para diferenciar cada módulo de la aplicación:')
    
    color_table = doc.add_table(rows=1, cols=3)
    color_table.style = 'Table Grid'
    hdr = color_table.rows[0].cells
    hdr[0].text = 'Color'
    hdr[1].text = 'Uso en la Interfaz'
    hdr[2].text = 'Código Hexadecimal'
    
    palette = [
        ('Naranja Rojizo', 'Login, Carrito y Botones', '#ff6347'),
        ('Verde Bosque', 'Módulo Rescate y Salud', '#2e7d32'),
        ('Azul Océano', 'Módulo de Recetas', '#0288d1'),
        ('Ámbar / Oro', 'Descuentos y Badges', '#f57c00'),
        ('Gris Fondo', 'Fondos de Pantalla', '#f5f5f5'),
        ('Blanco Puro', 'Tarjetas y NavBars', '#ffffff')
    ]
    for n, u, h in palette:
        row = color_table.add_row().cells
        row[0].text = n
        row[1].text = u
        row[2].text = h

    # 5. DIAGRAMA DE BASE DE DATOS (REQUISITO)
    doc.add_heading('5. DIAGRAMA Y LÓGICA DE BASE DE DATOS', level=1)
    doc.add_paragraph('La aplicación implementa un sistema híbrido de persistencia de datos:')
    
    doc.add_paragraph('Persistencia Local (AsyncStorage): Almacena de forma permanente los usuarios registrados, la sesión activa y el estado del carrito de compras, funcionando como una base de datos clave-valor en el dispositivo.', style='List Bullet')
    doc.add_paragraph('Consumo de API (TheMealDB): Suministra la base de datos de platillos, ingredientes e instrucciones en tiempo real mediante peticiones REST JSON.', style='List Bullet')

    doc.add_page_break()

    # 6. Tecnologías y Navegación
    doc.add_heading('6. TECNOLOGÍAS Y NAVEGACIÓN', level=1)
    doc.add_paragraph('Navegación Stack (4 niveles): Login > Main > Details > Checkout.')
    doc.add_paragraph('Navegación Tab (5 secciones): Rescate, Dietas, Recetas, Carrito, Perfil.')
    doc.add_paragraph('Frontend: React Native Paper con UI adaptable.')

    # 7. APK
    doc.add_heading('7. PROCESO DE CREACIÓN DE APK', level=1)
    doc.add_paragraph('Se utilizó EAS (Expo Application Services) para compilar el archivo ejecutable nativo para Android (APK). El comando principal utilizado fue:')
    doc.add_paragraph('eas build -p android --profile preview', style='Intense Quote')

    # Conclusión
    doc.add_heading('8. CONCLUSIÓN', level=1)
    doc.add_paragraph('FastBite cumple con la totalidad de los requisitos técnicos de las Unidades 3 y 4 de Desarrollo Móvil, integrando componentes avanzados de React Native y servicios externos de traducción.')

    doc.save('Manual_FastBite_ENTREGA.docx')
    print("Manual de ENTREGA generado con éxito.")

if __name__ == "__main__":
    create_manual()
