from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()

    # Portada
    title = doc.add_heading('MANUAL DE PROCESO DE CREACIÓN — PROYECTO FASTBITE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('\n' * 3)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('DESARROLLO DE APLICACIONES MÓVILES\nUNIDADES 3 Y 4\n\nALUMNO: Henrique\nMATRÍCULA: 2302070\nFECHA: 14 de Abril de 2026')
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph('\n' * 2)
    repo = doc.add_paragraph()
    repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    repo.add_run('REPOSITORIO GITHUB:\nhttps://github.com/2302070-commits/Proyecto_Unidad_4_FastBite').font.color.rgb = RGBColor(0, 0, 255)

    doc.add_page_break()

    # 1. Introducción
    doc.add_heading('1. INTRODUCCIÓN', level=1)
    doc.add_paragraph('FastBite es una solución móvil de vanguardia construida con React Native y Expo. El objetivo primordial es fusionar la responsabilidad social (rescate de alimentos) con la salud personal (dietas especializadas) y la educación culinaria (recetario inteligente).')

    # 2. Problemática
    doc.add_heading('2. PROBLEMÁTICA DETECTADA', level=1)
    doc.add_paragraph('Problemática 1 — Desperdicio: México pierde más de 20 millones de toneladas de comida al año mientras pequeños negocios locales carecen de canales para vender excedentes.', style='List Bullet')
    doc.add_paragraph('Problemática 2 — Dietas: Existe una brecha de información para usuarios con necesidades alimentarias específicas (veganos/vegetarianos) que requieren datos precisos y en español.', style='List Bullet')

    # 3. Propuesta de Solución
    doc.add_heading('3. PROPUESTA DE SOLUCIÓN', level=1)
    doc.add_paragraph('La aplicación se divide en tres ejes estratégicos:')
    
    p1 = doc.add_paragraph()
    p1.add_run('🌱 Módulo de Rescate:').bold = True
    p1.add_run(' Permite comprar platillos locales con descuentos del 50-70% para evitar el desperdicio.')
    
    p2 = doc.add_paragraph()
    p2.add_run('❤️ Módulo de Dietas:').bold = True
    p2.add_run(' Filtra platillos saludables y especializados con conteo calórico estimado.')
    
    p3 = doc.add_paragraph()
    p3.add_run('📖 Módulo de Recetas:').bold = True
    p3.add_run(' Buscador con inteligencia de traducción que conecta con APIs internacionales.')

    # 4. Tecnologías
    doc.add_heading('4. TECNOLOGÍAS Y LIBRERÍAS', level=1)
    tech_table = doc.add_table(rows=1, cols=2)
    tech_table.style = 'Table Grid'
    hdr = tech_table.rows[0].cells
    hdr[0].text = 'Tecnología / Librería'
    hdr[1].text = 'Propósito'
    
    libs = [
        ('React Native 0.81.5', 'Framework principal'),
        ('Expo SDK 54', 'Entorno de desarrollo y compilación'),
        ('React Navigation (Stack/Tabs)', 'Gestión de rutas y pantallas'),
        ('React Native Paper', 'Sistema de componentes Material Design'),
        ('AsyncStorage', 'Persistencia de datos (Base de Datos Local)'),
        ('TheMealDB API', 'Fuente de datos externa interactiva'),
        ('Google Translate API', 'Motor de traducción en tiempo real')
    ]
    for lib, purpose in libs:
        row = tech_table.add_row().cells
        row[0].text = lib
        row[1].text = purpose

    # 5. Navegación
    doc.add_heading('5. ESTRUCTURA DE NAVEGACIÓN', level=1)
    doc.add_paragraph('Para cumplir con los requisitos académicos, se implementó una arquitectura de navegación anidada:')
    doc.add_paragraph('Stack Navigator principal con 4 pantallas (Login, Main, Detail, Checkout).', style='List Bullet')
    doc.add_paragraph('Tab Navigator con 5 secciones activas (Rescate, Dietas, Recetas, Carrito, Perfil).', style='List Bullet')

    # 6. Funcionalidades de Desarrollo
    doc.add_heading('6. FUNCIONALIDADES TÉCNICAS IMPLEMENTADAS', level=1)
    doc.add_paragraph('Se aplicaron los siguientes conceptos de programación móvil:')
    
    func_table = doc.add_table(rows=1, cols=2)
    func_table.style = 'Table Grid'
    h = func_table.rows[0].cells
    h[0].text = 'Concepto'
    h[1].text = 'Aplicación en el Proyecto'
    
    funcs = [
        ('Hooks (useState, useEffect)', 'Manejo de estados de búsqueda y carga de API.'),
        ('Context API', 'Carro de compras y sesión de usuario global.'),
        ('Props y Params', 'Pase de datos dinámicos entre listas y detalles.'),
        ('Fetch / API Rest', 'Consumo de servicios externos asíncronos.'),
        ('AsyncStorage', 'Guardado persistente de la sesión y el carrito.')
    ]
    for c, a in funcs:
        r = func_table.add_row().cells
        r[0].text = c
        r[1].text = a

    # 7. Proceso de APK
    doc.add_heading('7. PROCESO DE COMPILACIÓN (APK)', level=1)
    doc.add_paragraph('La generación del ejecutable se realizó mediante los servidores de EAS de Expo. El proceso incluyó la configuración del slug, el package name (com.henry.fastbite) y la generación del Android Keystore en la nube.')

    # Conclusión
    doc.add_page_break()
    doc.add_heading('CONCLUSIÓN', level=1)
    doc.add_paragraph('La aplicación FastBite no solo cumple con los requisitos técnicos de las Unidades 3 y 4, sino que propone un modelo de negocio socialmente responsable. El uso de tecnologías modernas asegura que la aplicación sea escalable, rápida y fácil de usar para el usuario final.')

    doc.save('Manual_Entrega_FastBite.docx')
    print("Documento completo guardado.")

if __name__ == "__main__":
    create_manual()
