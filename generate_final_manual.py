from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()

    # Portada Profesional
    title = doc.add_heading('MANUAL TÉCNICO DE DESARROLLO — FASTBITE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n' * 3)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('DESARROLLO DE APLICACIONES MÓVILES\nUNIDADES 3 Y 4\n\nALUMNO: Henrique\nMATRÍCULA: 2302070\nUNIVERSIDAD / INSTITUCIÓN\n14 de Abril de 2026')
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph('\n' * 2)
    repo = doc.add_paragraph()
    repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    repo.add_run('REPOSITORIO GITHUB:\nhttps://github.com/2302070-commits/Proyecto_Unidad_4_FastBite').font.color.rgb = RGBColor(0, 0, 255)

    doc.add_page_break()

    # Secciones Técnicas
    sections = [
        ('1. INTRODUCCIÓN', 'FastBite es una aplicación híbrida de alto rendimiento diseñada para abordar problemáticas sociales críticas como el desperdicio de alimentos y el acceso a nutrición especializada. El sistema integra tecnologías geo-locales simuladas y APIs de datos internacionales para ofrecer una experiencia de usuario fluida y educativa.'),
        
        ('2. PROBLEMÁTICA Y JUSTIFICACIÓN', 'El proyecto surge ante la alarmante cifra de desperdicio alimentario en México y la falta de aplicaciones accesibles que traduzcan información nutricional compleja al español de manera automática.'),
        
        ('3. ARQUITECTURA DEL SISTEMA', 'Se utiliza una arquitectura limpia basada en "Folders" para separar la lógica de negocio, los componentes visuales y la navegación. El manejo del estado global se delegó a la Context API de React, eliminando el "prop-drilling" y mejorando el rendimiento.')
    ]

    for title_text, body_text in sections:
        doc.add_heading(title_text, level=1)
        doc.add_paragraph(body_text)

    # Tabla de Navegación
    doc.add_heading('4. ESTRUCTURA DE NAVEGACIÓN (UNIDADES 3 Y 4)', level=1)
    nav_table = doc.add_table(rows=1, cols=2)
    nav_table.style = 'Table Grid'
    h = nav_table.rows[0].cells
    h[0].text = 'Tipo de Navegador'
    h[1].text = 'Pantallas / Secciones'
    
    navs = [
        ('Stack Navigator (Principal)', 'Login, MainTabs, FoodDetail, Checkout'),
        ('Tab Navigator (Secciones)', 'Rescate, Dietas, Recetas, Carrito, Perfil')
    ]
    for t, s in navs:
        r = nav_table.add_row().cells
        r[0].text = t
        r[1].text = s

    # Tabla de Librerías
    doc.add_heading('5. LIBRERÍAS Y BIBLIOGRAFÍAS', level=1)
    lib_table = doc.add_table(rows=1, cols=2)
    lib_table.style = 'Table Grid'
    h = lib_table.rows[0].cells
    h[0].text = 'Librería'
    h[1].text = 'Funcionalidad'
    
    libs = [
        ('Expo SDK 54', 'Gestión de Build y entorno nativo.'),
        ('React Navigation 7', 'Control de rutas y stacks.'),
        ('React Native Paper', 'Diseño de UI Material Design 3.'),
        ('AsyncStorage', 'Persistencia de datos (BD Local).'),
        ('TheMealDB API', 'Suministro de datos de platillos.'),
        ('Google Translate', 'Traducción dinámica de instrucciones.')
    ]
    for l, f in libs:
        r = lib_table.add_row().cells
        r[0].text = l
        r[1].text = f

    # Tecnologías de Programación
    doc.add_heading('6. CONCEPTOS DE PROGRAMACIÓN APLICADOS', level=1)
    doc.add_paragraph('Variables de estado: useState para control de carritos y búsquedas.', style='List Bullet')
    doc.add_paragraph('Variables de propiedades: Props para pasar ID de comida a la pantalla de detalle.', style='List Bullet')
    doc.add_paragraph('Hooks de efecto: useEffect para llamadas a APIs y persistencia.', style='List Bullet')
    doc.add_paragraph('Estado Global: useContext para Auth y Carrito.', style='List Bullet')

    # Compilación
    doc.add_heading('7. GENERACIÓN DEL EJECUTABLE (APK)', level=1)
    doc.add_paragraph('El entregable incluye un archivo APK compilado en la nube mediante Expo Application Services (EAS). Este archivo permite la instalación directa en dispositivos Android reales sin necesidad de cables.')

    doc.add_page_break()
    doc.add_heading('CONCLUSIÓN FINAL', level=1)
    doc.add_paragraph('El proyecto FastBite representa la integración exitosa de los conocimientos adquiridos en el curso, demostrando competencia en consumo de APIs, manejo de estados, persistencia de datos y diseño responsive en plataformas móviles.')

    doc.save('Manual_FastBite_FINAL.docx')
    print("Manual FINAL generado con éxito.")

if __name__ == "__main__":
    create_manual()
